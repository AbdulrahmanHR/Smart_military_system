"""
Document processing module for military training documents
"""
import os
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Tuple

import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

from config.settings import config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document ingestion, processing, and chunking for military training materials with Arabic support"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Arabic category keywords mapping
        self.arabic_category_keywords = {
            "Tactical Procedures": {
                "en": ['tactical', 'combat', 'strategy', 'maneuver', 'tactics'],
                "ar": ['تكتيكي', 'قتال', 'استراتيجية', 'مناورة', 'تكتيكات', 'خطة', 'هجوم', 'دفاع']
            },
            "Equipment Training": {
                "en": ['equipment', 'weapon', 'gear', 'maintenance', 'tools'],
                "ar": ['معدات', 'سلاح', 'أدوات', 'صيانة', 'تجهيزات', 'آلات', 'أسلحة']
            },
            "Emergency Protocols": {
                "en": ['emergency', 'crisis', 'evacuation', 'medical', 'urgent'],
                "ar": ['طوارئ', 'أزمة', 'إخلاء', 'طبي', 'عاجل', 'إسعاف', 'حالة طارئة']
            },
            "Leadership & Coordination": {
                "en": ['leadership', 'command', 'coordination', 'team', 'management'],
                "ar": ['قيادة', 'قائد', 'تنسيق', 'فريق', 'إدارة', 'تعاون', 'ضابط']
            },
            "Physical Training": {
                "en": ['physical', 'fitness', 'training', 'exercise', 'conditioning'],
                "ar": ['بدني', 'لياقة', 'تدريب', 'تمرين', 'رياضة', 'تكييف', 'قوة']
            },
            "Safety Procedures": {
                "en": ['safety', 'security', 'protection', 'risk', 'precaution'],
                "ar": ['أمان', 'أمن', 'حماية', 'خطر', 'احتياط', 'وقاية', 'سلامة']
            },
            "Communication Protocols": {
                "en": ['communication', 'radio', 'signal', 'intel', 'transmission'],
                "ar": ['اتصال', 'راديو', 'إشارة', 'استخبارات', 'إرسال', 'لاسلكي', 'تواصل']
            },
            "Mission Planning": {
                "en": ['mission', 'planning', 'operation', 'briefing', 'objective'],
                "ar": ['مهمة', 'تخطيط', 'عملية', 'إحاطة', 'هدف', 'خطة عمل', 'مهام']
            }
        }
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text files with enhanced Arabic support"""
        try:
            # Try UTF-8 first (preferred for Arabic)
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Fallback to other encodings that might contain Arabic
            for encoding in ['utf-16', 'cp1256', 'iso-8859-6']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        logger.info(f"Successfully read {file_path} with encoding: {encoding}")
                        return content
                except (UnicodeDecodeError, UnicodeError):
                    continue
            logger.error(f"Could not decode {file_path} with any supported encoding")
            return ""
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def detect_language(self, text: str) -> str:
        """Simple language detection for Arabic vs English text"""
        if not text:
            return "en"
        
        # Count Arabic characters
        arabic_chars = len(re.findall(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        total_chars = len(text.strip())
        
        if total_chars == 0:
            return "en"
        
        arabic_ratio = arabic_chars / total_chars
        english_ratio = english_chars / total_chars
        
        # If more than 10% Arabic characters, consider it Arabic
        if arabic_ratio > 0.1 and arabic_ratio > english_ratio:
            return "ar"
        else:
            return "en"
    
    def determine_document_category(self, filename: str, content: str) -> str:
        """Determine document category based on filename and content with Arabic support"""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        # Detect primary language of content
        detected_lang = self.detect_language(content)
        
        # Score each category based on keyword matches
        category_scores = {}
        
        for category, keywords in self.arabic_category_keywords.items():
            score = 0
            
            # Check English keywords
            for keyword in keywords["en"]:
                if keyword in filename_lower:
                    score += 2  # Filename matches get higher weight
                if keyword in content_lower:
                    score += 1
            
            # Check Arabic keywords
            for keyword in keywords["ar"]:
                if keyword in filename:  # Don't lowercase Arabic text
                    score += 2
                if keyword in content:
                    score += 1
            
            category_scores[category] = score
        
        # Find category with highest score
        if category_scores:
            best_category = max(category_scores.items(), key=lambda x: x[1])
            if best_category[1] > 0:
                return best_category[0]
        
        return "General Training"
    
    def process_document(self, file_path: Path) -> List[LangchainDocument]:
        """Process a single document and return chunked documents"""
        logger.info(f"Processing document: {file_path}")
        
        # Extract text based on file type
        file_extension = file_path.suffix.lower()
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return []
        
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return []
        
        # Determine document category and language
        category = self.determine_document_category(file_path.name, text)
        detected_language = self.detect_language(text)
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangChain documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "category": category,
                "language": detected_language,
                "chunk_index": i,
                "total_chunks": len(chunks)
            }
            documents.append(LangchainDocument(page_content=chunk, metadata=metadata))
        
        logger.info(f"Created {len(documents)} chunks from {file_path}")
        return documents
    
    def process_directory(self, directory_path: Path) -> List[LangchainDocument]:
        """Process all supported documents in a directory"""
        supported_extensions = {'.pdf', '.docx', '.txt'}
        all_documents = []
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                documents = self.process_document(file_path)
                all_documents.extend(documents)
        
        logger.info(f"Processed {len(all_documents)} total chunks from {directory_path}")
        return all_documents
    
    def get_document_stats(self, documents: List[LangchainDocument]) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not documents:
            return {"total_chunks": 0, "categories": {}, "sources": []}
        
        categories = {}
        sources = set()
        
        for doc in documents:
            category = doc.metadata.get('category', 'Unknown')
            categories[category] = categories.get(category, 0) + 1
            sources.add(doc.metadata.get('filename', 'Unknown'))
        
        return {
            "total_chunks": len(documents),
            "categories": categories,
            "sources": list(sources)
        }
